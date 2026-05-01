"""
combine_docx.py — Merge dt_docs2/02_ABSTRACT through 10_APPENDICES into one DOCX.
Output: dt_docs2/DISSERTATION_COMPLETE.docx
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

BASE = Path(__file__).parent
SRC  = BASE / "dt_docs2"
OUT  = SRC / "DISSERTATION_COMPLETE.docx"

# Files in order: 02 through 10
files = sorted([
    f for f in SRC.glob("*.docx")
    if f.stem[:2].isdigit() and 2 <= int(f.stem[:2]) <= 10
])

print(f"\nMerging {len(files)} files:\n")
for f in files:
    print(f"  {f.name}")


def add_page_break(doc):
    """Insert a page break at end of document."""
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# Open first file as master
master   = Document(str(files[0]))
composer = Composer(master)

for f in files[1:]:
    doc = Document(str(f))
    # Add page break before appending next chapter
    add_page_break(master)
    composer.append(doc)

composer.save(str(OUT))
print(f"\nSaved -> {OUT}")

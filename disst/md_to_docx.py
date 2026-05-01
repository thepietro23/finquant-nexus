"""
Convert all dissertation MD files → DOCX with RRU formatting.
Output folder: fqn1/disst/dt_docs/
Font: Times New Roman | Body: 12pt | H1: 14pt Bold | H2: 12pt Bold
Margins: Left 38mm, Right/Top/Bottom 25.4mm | Line spacing: 1.5
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SRC = Path(r"e:\Sem4\Clg_Project\finquant-nexus\disst\chap_content")
DST = Path(r"e:\Sem4\Clg_Project\finquant-nexus\disst\dt_docs")
DST.mkdir(parents=True, exist_ok=True)

FONT = "Times New Roman"

MD_FILES = sorted(SRC.glob("*.md"))


def set_doc_defaults(doc):
    """Apply RRU page setup to document."""
    sec = doc.sections[0]
    sec.page_width  = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin   = Mm(38)
    sec.right_margin  = Mm(25.4)
    sec.top_margin    = Mm(25.4)
    sec.bottom_margin = Mm(25.4)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name   = FONT
    style.font.size   = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after       = Pt(6)
    pf.space_before      = Pt(0)


def fmt_run(run, bold=False, italic=False, size=12, color=None):
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_spacing(para, before=0, after=6, line=WD_LINE_SPACING.ONE_POINT_FIVE):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = line


def apply_inline(para, text, base_size=12, base_bold=False):
    """Apply **bold** inline markdown to a paragraph."""
    # Strip HTML tags but keep inner text
    text = re.sub(r'<[^>]+>', '', text)
    # Strip markdown links [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Handle **bold** and *italic* inline
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            fmt_run(run, bold=True, size=base_size)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            fmt_run(run, italic=True, size=base_size)
        else:
            run = para.add_run(part)
            fmt_run(run, bold=base_bold, size=base_size)


def add_heading(doc, text, level):
    """Add heading with RRU formatting."""
    para = doc.add_paragraph()
    set_para_spacing(para, before=12 if level == 1 else 8, after=4)
    size = 14 if level == 1 else 12
    apply_inline(para, text, base_size=size, base_bold=True)
    return para


def add_table_from_md(doc, table_lines):
    """Convert markdown table lines to a DOCX table."""
    # Filter out separator lines (---|---|---)
    data_rows = [l for l in table_lines if not re.match(r'^\|[-:| ]+\|$', l.strip())]
    if not data_rows:
        return

    rows = []
    for line in data_rows:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    # Pad rows
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'

    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            para = cell.paragraphs[0]
            is_header = (ri == 0)
            apply_inline(para, re.sub(r'<[^>]+>', '', cell_text), base_size=11, base_bold=is_header)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()  # spacing after table


def add_code_block(doc, lines):
    """Add code/yaml block as monospace."""
    para = doc.add_paragraph()
    set_para_spacing(para, before=4, after=4)
    run = para.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light shading
    try:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F2F2F2')
        para._p.get_or_add_pPr().append(shading)
    except Exception:
        pass


def _math_unicode(expr: str) -> str:
    """Convert ASCII math notation to Unicode for display equations."""
    # Greek letters
    expr = re.sub(r'\bmu\b',      'μ',  expr)
    expr = re.sub(r'\balpha\b',   'α',  expr)
    expr = re.sub(r'\bbeta\b',    'β',  expr)
    expr = re.sub(r'\bgamma\b',   'γ',  expr)
    expr = re.sub(r'\bdelta\b',   'δ',  expr)
    expr = re.sub(r'\bsigma\b',   'σ',  expr)
    expr = re.sub(r'\btheta\b',   'θ',  expr)
    expr = re.sub(r'\blambda\b',  'λ',  expr)
    expr = re.sub(r'\bepsilon\b', 'ε',  expr)
    expr = re.sub(r'\bphi\b',     'φ',  expr)
    expr = re.sub(r'\bpi\b',      'π',  expr)
    expr = re.sub(r'\bomega\b',   'ω',  expr)
    # Superscripts
    expr = expr.replace('^2',  '²')
    expr = expr.replace('^3',  '³')
    expr = expr.replace('^T',  'ᵀ')
    expr = expr.replace('^-1', '⁻¹')
    # Subscripts (common cases)
    expr = re.sub(r'_\{([^}]+)\}', lambda m: ''.join(
        '₀₁₂₃₄₅₆₇₈₉'[int(c)] if c.isdigit() else c for c in m.group(1)), expr)
    expr = re.sub(r'_([0-9])', lambda m: '₀₁₂₃₄₅₆₇₈₉'[int(m.group(1))], expr)
    # Operators
    expr = expr.replace('||', '‖')
    expr = expr.replace('!=', '≠')
    expr = expr.replace('>=', '≥')
    expr = expr.replace('<=', '≤')
    expr = expr.replace('->',  '→')
    expr = expr.replace('<-',  '←')
    expr = re.sub(r'\bsqrt\b', '√', expr)
    expr = re.sub(r'\bsum\b',  'Σ', expr)
    expr = re.sub(r'\bprod\b', 'Π', expr)
    expr = re.sub(r'\binf\b',  '∞', expr)
    # N(0, ...) stays as is — common enough to be readable
    return expr


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    set_doc_defaults(doc)

    lines = md_path.read_text(encoding='utf-8', errors='replace').splitlines()

    i = 0
    table_buffer = []
    code_buffer  = []
    in_code = False
    in_table = False

    def flush_table():
        nonlocal table_buffer, in_table
        if table_buffer:
            add_table_from_md(doc, table_buffer)
        table_buffer = []
        in_table = False

    def flush_code():
        nonlocal code_buffer, in_code
        if code_buffer:
            add_code_block(doc, code_buffer)
        code_buffer = []
        in_code = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # --- Code block toggle ---
        if line.strip().startswith('```'):
            if in_code:
                flush_code()
            else:
                if in_table:
                    flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # --- Table detection ---
        is_table_line = bool(line.strip().startswith('|') and '|' in line[1:])
        if is_table_line:
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        stripped = line.strip()

        # Skip empty lines (paragraph break already handled by spacing)
        if not stripped:
            i += 1
            continue

        # --- Horizontal rule --- skip entirely, don't render
        if re.match(r'^---+$', stripped) or re.match(r'^\*\*\*+$', stripped):
            i += 1
            continue

        # --- Headings ---
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            # Skip file meta title: # Chapter X — Title (top of each file)
            if level == 1 and re.search(r'Chapter \d+', text):
                i += 1
                continue
            # Skip ALL metadata lines including Reference: lines in headings
            if re.search(r'Target:|Status:|Word count:|Reference:.*DISSERTATION|DISSERTATION_FORMATTING', text):
                i += 1
                continue
            add_heading(doc, text, min(level, 4))
            i += 1
            continue

        # --- Blockquote / note lines ---
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            # Skip ALL metadata/formatting instruction lines
            if any(kw in text for kw in [
                'Reference:', 'Writing rules:', 'prompt.md', 'Last updated:',
                'DISSERTATION_FORMATTING', 'Arabic page numbering', 'page numbering from',
                'Annexure', 'update after all figures', 'update after all tables',
                'update actual page numbers'
            ]):
                i += 1
                continue
            if text:
                para = doc.add_paragraph()
                set_para_spacing(para, before=2, after=2)
                run = para.add_run(text)
                run.font.name   = FONT
                run.font.size   = Pt(11)
                run.font.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue

        # --- Bullet list ---
        m_bullet = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if m_bullet:
            indent = len(m_bullet.group(1))
            text   = m_bullet.group(3)
            para   = doc.add_paragraph(style='List Bullet')
            set_para_spacing(para, before=1, after=1)
            para.paragraph_format.left_indent = Mm(5 + indent * 3)
            apply_inline(para, text, base_size=12)
            i += 1
            continue

        # --- Numbered list ---
        m_num = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m_num:
            text = m_num.group(3)
            para = doc.add_paragraph(style='List Number')
            set_para_spacing(para, before=1, after=1)
            apply_inline(para, text, base_size=12)
            i += 1
            continue

        # --- Plain paragraph ---
        # Skip footer metadata lines
        if re.match(r'^\*Reference:.*\*$', stripped) or re.match(r'^\*Last updated:.*\*$', stripped):
            i += 1
            continue

        # --- Math equation block (4-space or tab indent = display equation) ---
        if raw.startswith('    ') or raw.startswith('\t'):
            eq = stripped
            eq = _math_unicode(eq)
            para = doc.add_paragraph()
            set_para_spacing(para, before=6, after=6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.left_indent  = Mm(10)
            para.paragraph_format.right_indent = Mm(10)
            run = para.add_run(eq)
            run.font.name   = 'Cambria Math'
            run.font.size   = Pt(12)
            run.font.italic = False
            i += 1
            continue

        # Strip HTML tags but keep the visible text (especially [INSERT ...] markers)
        clean = re.sub(r'<span[^>]*>', '', stripped)
        clean = re.sub(r'</span>', '', clean)
        clean = _math_unicode(clean.strip())
        if clean:
            para = doc.add_paragraph()
            set_para_spacing(para, before=0, after=6)
            apply_inline(para, clean, base_size=12)

        i += 1

    # Flush any remaining buffers
    if in_table:
        flush_table()
    if in_code:
        flush_code()

    doc.save(str(docx_path))
    print(f"  OK {docx_path.name}")


# ── Main ──────────────────────────────────────────────────────────────────────
print(f"\nConverting {len(MD_FILES)} files to {DST}\n")
ok_count = 0
skipped = []
for md in MD_FILES:
    docx_name = md.stem + ".docx"
    try:
        md_to_docx(md, DST / docx_name)
        ok_count += 1
    except PermissionError:
        print(f"  SKIP {docx_name} (file is open in MS Word)")
        skipped.append(docx_name)
    except Exception as e:
        print(f"  FAIL {docx_name}: {e}")
        skipped.append(docx_name)

print(f"\nDone — {ok_count} of {len(MD_FILES)} files written to dt_docs/")
if skipped:
    print(f"Skipped: {', '.join(skipped)}")
